"""
export_service.py — Excel (openpyxl) and Word (python-docx) exports
                    with full Arabic RTL support.
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

# ─────────────────────────────────────────────────────────────
# EXCEL EXPORT — Risk Register
# ─────────────────────────────────────────────────────────────

def export_risk_register_xlsx(risks: List[Any], project_name: str = "المشروع") -> bytes:
    """
    Export risk register items to Excel with Arabic headers and RTL formatting.
    Returns bytes of the .xlsx file.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import (
            Font, PatternFill, Alignment, Border, Side, GradientFill
        )
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise RuntimeError("openpyxl not installed. Run: pip install openpyxl")

    wb = Workbook()
    ws = wb.active
    ws.title = "سجل المخاطر"

    # RTL sheet direction
    ws.sheet_view.rightToLeft = True

    # Color palette
    HEADER_FILL = PatternFill("solid", fgColor="1B4F72")
    SUBHEADER_FILL = PatternFill("solid", fgColor="2E86AB")
    HIGH_FILL = PatternFill("solid", fgColor="FDEDEC")
    MED_FILL = PatternFill("solid", fgColor="FEF9E7")
    LOW_FILL = PatternFill("solid", fgColor="EAFAF1")
    OPP_FILL = PatternFill("solid", fgColor="EBF5FB")
    ALT_FILL = PatternFill("solid", fgColor="F5F7FA")

    thin = Side(border_style="thin", color="DDE3EC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def hdr_cell(ws, row, col, value, width=18):
        c = ws.cell(row=row, column=col, value=value)
        c.font = Font(name="Cairo", bold=True, color="FFFFFF", size=11)
        c.fill = HEADER_FILL
        c.alignment = Alignment(horizontal="center", vertical="center",
                                wrap_text=True, reading_order=2)
        c.border = border
        ws.column_dimensions[get_column_letter(col)].width = width
        return c

    def data_cell(ws, row, col, value, fill=None, bold=False, color="1A2535"):
        c = ws.cell(row=row, column=col, value=value)
        c.font = Font(name="Cairo", size=10, bold=bold, color=color)
        if fill:
            c.fill = fill
        c.alignment = Alignment(horizontal="right", vertical="center",
                                wrap_text=True, reading_order=2)
        c.border = border
        return c

    # Title row
    ws.merge_cells("A1:U1")
    title_cell = ws["A1"]
    title_cell.value = f"سجل المخاطر — {project_name}  |  تاريخ التصدير: {datetime.now().strftime('%Y-%m-%d')}"
    title_cell.font = Font(name="Cairo", bold=True, size=14, color="FFFFFF")
    title_cell.fill = PatternFill("solid", fgColor="0D2F47")
    title_cell.alignment = Alignment(horizontal="center", vertical="center",
                                     reading_order=2)
    ws.row_dimensions[1].height = 35

    # Headers row 2
    headers = [
        ("رقم المخاطرة", 14), ("المستوى", 12), ("النوع", 10),
        ("الفئة", 18), ("الفئة الفرعية", 16),
        ("السبب", 22), ("الحدث", 22), ("التأثير", 22),
        ("جملة س→ح→ت", 28), ("المحفز", 16),
        ("الاحتمالية", 12), ("ت. التكلفة", 12), ("ت. الوقت", 12),
        ("ت. النطاق", 12), ("ت. الجودة", 12),
        ("التأثير المركب", 14), ("الدرجة", 10), ("الأولوية", 12),
        ("استراتيجية الاستجابة", 20), ("مالك المخاطرة", 18),
        ("حالة الموافقة", 16),
    ]
    for col_idx, (header, width) in enumerate(headers, start=1):
        hdr_cell(ws, 2, col_idx, header, width)
    ws.row_dimensions[2].height = 40

    # Freeze top 2 rows
    ws.freeze_panes = "A3"

    # Priority → fill
    priority_fills = {
        "حرجة": HIGH_FILL,
        "عالية": PatternFill("solid", fgColor="FDEBD0"),
        "متوسطة": MED_FILL,
        "منخفضة": LOW_FILL,
    }

    # Data rows
    for row_num, risk in enumerate(risks, start=3):
        row_fill = ALT_FILL if row_num % 2 == 0 else None
        priority = getattr(risk, "priority", "منخفضة")
        score = getattr(risk, "score", 0)
        risk_type = getattr(risk, "risk_type", "")
        cell_fill = OPP_FILL if str(risk_type) == "opportunity" else priority_fills.get(priority, row_fill)

        row_data = [
            getattr(risk, "risk_id", ""),
            str(getattr(risk, "level", "")),
            "فرصة" if str(risk_type) == "opportunity" else "تهديد",
            getattr(risk, "category", ""),
            getattr(risk, "subcategory", "") or "",
            getattr(risk, "cause", "") or "",
            getattr(risk, "event", "") or "",
            getattr(risk, "impact_description", "") or "",
            getattr(risk, "cei_statement", "") or "",
            getattr(risk, "trigger", "") or "",
            getattr(risk, "probability", 0),
            getattr(risk, "cost_impact", 0),
            getattr(risk, "time_impact", 0),
            getattr(risk, "scope_impact", 0),
            getattr(risk, "quality_impact", 0),
            getattr(risk, "composite_impact", 0),
            score,
            priority,
            str(getattr(risk, "response_strategy", "") or ""),
            getattr(risk, "risk_owner", "") or "",
            str(getattr(risk, "approval_status", "") or ""),
        ]
        for col_idx, value in enumerate(row_data, start=1):
            c = data_cell(ws, row_num, col_idx, value, fill=cell_fill)
            # Score coloring
            if col_idx == 17:  # Score column
                if isinstance(score, (int, float)):
                    if score >= 15:
                        c.font = Font(name="Cairo", size=10, bold=True, color="C0392B")
                    elif score >= 9:
                        c.font = Font(name="Cairo", size=10, bold=True, color="D35400")
        ws.row_dimensions[row_num].height = 30

    # Add auto-filter
    ws.auto_filter.ref = f"A2:U{len(risks) + 2}"

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


# ─────────────────────────────────────────────────────────────
# EXCEL EXPORT — Response Tracking
# ─────────────────────────────────────────────────────────────

def export_tracking_xlsx(items: List[Any], project_name: str = "المشروع") -> bytes:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise RuntimeError("openpyxl not installed")

    wb = Workbook()
    ws = wb.active
    ws.title = "متابعة الاستجابة"
    ws.sheet_view.rightToLeft = True

    thin = Side(border_style="thin", color="DDE3EC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    HEADER_FILL = PatternFill("solid", fgColor="1B4F72")

    # Title
    ws.merge_cells("A1:L1")
    ws["A1"].value = f"سجل متابعة الاستجابة — {project_name}  |  {datetime.now().strftime('%Y-%m-%d')}"
    ws["A1"].font = Font(name="Cairo", bold=True, size=13, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor="0D2F47")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center", reading_order=2)
    ws.row_dimensions[1].height = 30

    headers = [
        "رقم المخاطرة", "إجراء الاستجابة", "المسؤول",
        "بداية مخططة", "نهاية مخططة", "نهاية فعلية",
        "التقدم %", "الحالة الحالية", "تصعيد مطلوب",
        "ملاحظة الدليل", "فعالية الاستجابة", "توصية الإغلاق",
    ]
    col_widths = [14, 30, 18, 14, 14, 14, 12, 14, 14, 22, 18, 22]
    for col_idx, (header, width) in enumerate(zip(headers, col_widths), start=1):
        c = ws.cell(row=2, column=col_idx, value=header)
        c.font = Font(name="Cairo", bold=True, color="FFFFFF", size=10)
        c.fill = HEADER_FILL
        c.alignment = Alignment(horizontal="center", vertical="center",
                                wrap_text=True, reading_order=2)
        c.border = border
        ws.column_dimensions[get_column_letter(col_idx)].width = width
    ws.row_dimensions[2].height = 36
    ws.freeze_panes = "A3"

    STATUS_COLORS = {
        "completed": "EAFAF1",
        "in_progress": "EBF5FB",
        "delayed": "FDEDEC",
        "cancelled": "F5F7FA",
        "not_started": "FEFEFE",
    }

    for row_num, item in enumerate(items, start=3):
        status_str = str(getattr(item, "current_status", "not_started"))
        fill_color = STATUS_COLORS.get(status_str, "FEFEFE")
        fill = PatternFill("solid", fgColor=fill_color)

        row_data = [
            getattr(item, "risk_id", ""),
            getattr(item, "response_action", ""),
            getattr(item, "action_owner", "") or "",
            getattr(item, "planned_start", "") or "",
            getattr(item, "planned_finish", "") or "",
            getattr(item, "actual_finish", "") or "",
            f"{getattr(item, 'progress_pct', 0)}%",
            _translate_status(status_str),
            "نعم" if getattr(item, "escalation_required", False) else "لا",
            getattr(item, "evidence_note", "") or "",
            _translate_effectiveness(str(getattr(item, "effectiveness_rating", ""))),
            getattr(item, "closure_recommendation", "") or "",
        ]
        for col_idx, value in enumerate(row_data, start=1):
            c = ws.cell(row=row_num, column=col_idx, value=value)
            c.font = Font(name="Cairo", size=10)
            c.fill = fill
            c.alignment = Alignment(horizontal="right", vertical="center",
                                    wrap_text=True, reading_order=2)
            c.border = border
        ws.row_dimensions[row_num].height = 28

    ws.auto_filter.ref = f"A2:L{len(items) + 2}"
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


# ─────────────────────────────────────────────────────────────
# EXCEL EXPORT — Admin Master Export
# ─────────────────────────────────────────────────────────────

def export_admin_master_xlsx(
    users: List[Any],
    projects: List[Any],
    risks: List[Any],
    tracking: List[Any],
) -> bytes:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise RuntimeError("openpyxl not installed")

    wb = Workbook()

    def make_sheet(wb, title, headers, rows, header_color="1B4F72"):
        ws = wb.create_sheet(title=title)
        ws.sheet_view.rightToLeft = True
        thin = Side(border_style="thin", color="DDE3EC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        HFILL = PatternFill("solid", fgColor=header_color)
        for col_idx, (header, width) in enumerate(headers, start=1):
            c = ws.cell(row=1, column=col_idx, value=header)
            c.font = Font(name="Cairo", bold=True, color="FFFFFF", size=10)
            c.fill = HFILL
            c.alignment = Alignment(horizontal="center", vertical="center",
                                    wrap_text=True, reading_order=2)
            c.border = border
            ws.column_dimensions[get_column_letter(col_idx)].width = width
        ws.freeze_panes = "A2"
        for row_num, row_data in enumerate(rows, start=2):
            for col_idx, value in enumerate(row_data, start=1):
                c = ws.cell(row=row_num, column=col_idx, value=value)
                c.font = Font(name="Cairo", size=10)
                c.alignment = Alignment(horizontal="right", vertical="center", reading_order=2)
                c.border = border
        ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{len(rows)+1}"
        return ws

    # Sheet 1: Users
    user_headers = [
        ("الاسم الكامل", 22), ("البريد الإلكتروني", 28), ("الجوال", 16),
        ("الشركة", 22), ("الدولة", 14), ("الدور", 10), ("الحالة", 18),
        ("تاريخ التسجيل", 16), ("انتهاء الاشتراك", 16),
    ]
    user_rows = [
        [
            u.full_name, u.email, u.mobile or "", u.company or "", u.country or "",
            str(u.role), str(u.status),
            u.created_at.strftime("%Y-%m-%d") if u.created_at else "",
            u.activation_expires_at.strftime("%Y-%m-%d") if u.activation_expires_at else "",
        ]
        for u in users
    ]
    make_sheet(wb, "المستخدمون", user_headers, user_rows)

    # Sheet 2: Projects
    proj_headers = [
        ("اسم المشروع", 28), ("النوع", 16), ("صاحب المشروع", 20),
        ("قيمة العقد", 14), ("العملة", 10), ("المدة (شهر)", 12),
        ("شهية المخاطر", 14), ("تاريخ الإنشاء", 16),
    ]
    proj_rows = [
        [
            p.name, p.project_type or "", str(p.owner_id),
            p.contract_value or 0, p.currency,
            p.duration_months or 0, p.risk_appetite,
            p.created_at.strftime("%Y-%m-%d") if p.created_at else "",
        ]
        for p in projects
    ]
    make_sheet(wb, "المشاريع", proj_headers, proj_rows, header_color="2E86AB")

    # Sheet 3: Risks
    risk_headers = [
        ("رقم المخاطرة", 14), ("المشروع", 14), ("الفئة", 18),
        ("النوع", 10), ("الدرجة", 10), ("الأولوية", 12),
        ("مالك المخاطرة", 18), ("حالة الموافقة", 18), ("حالة الدورة", 14),
        ("تاريخ الإنشاء", 16),
    ]
    risk_rows = [
        [
            r.risk_id, str(r.project_id), r.category, str(r.risk_type),
            r.score, r.priority, r.risk_owner or "",
            str(r.approval_status), str(r.lifecycle_status),
            r.created_at.strftime("%Y-%m-%d") if r.created_at else "",
        ]
        for r in risks
    ]
    make_sheet(wb, "سجل المخاطر", risk_headers, risk_rows, header_color="C0392B")

    # Sheet 4: Tracking
    track_headers = [
        ("رقم المخاطرة", 12), ("المسؤول", 20), ("التقدم %", 12),
        ("الحالة", 16), ("تصعيد", 12), ("الفعالية", 16), ("تاريخ التحديث", 16),
    ]
    track_rows = [
        [
            str(t.risk_id), t.action_owner or "", t.progress_pct,
            _translate_status(str(t.current_status)),
            "نعم" if t.escalation_required else "لا",
            _translate_effectiveness(str(t.effectiveness_rating)),
            t.updated_at.strftime("%Y-%m-%d") if t.updated_at else "",
        ]
        for t in tracking
    ]
    make_sheet(wb, "متابعة الاستجابة", track_headers, track_rows, header_color="1E8449")

    # Remove default empty sheet
    if "Sheet" in wb.sheetnames:
        del wb["Sheet"]

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


# ─────────────────────────────────────────────────────────────
# WORD EXPORT — Risk Management Plan
# ─────────────────────────────────────────────────────────────

def export_risk_plan_docx(plan: Any, project: Any) -> bytes:
    """
    Export Risk Management Plan to Word with Arabic RTL formatting.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # Set RTL for entire document
    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "right")
        pPr.append(jc)

    def add_heading(doc, text, level=1):
        para = doc.add_heading(text, level=level)
        set_rtl(para)
        run = para.runs[0] if para.runs else para.add_run(text)
        run.font.name = "Cairo"
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
        return para

    def add_body(doc, text, bold=False):
        para = doc.add_paragraph(text)
        set_rtl(para)
        for run in para.runs:
            run.font.name = "Cairo"
            run.font.size = Pt(12)
            run.font.bold = bold
        return para

    def add_table_row(table, cells, header=False):
        row = table.add_row()
        for idx, text in enumerate(cells):
            cell = row.cells[idx]
            cell.text = str(text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Cairo"
                    run.font.size = Pt(10)
                    if header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        return row

    # ── Title Page ──
    title = doc.add_heading("خطة إدارة المخاطر", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Cairo"
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    doc.add_paragraph("")
    sub = doc.add_paragraph(getattr(project, "name", "المشروع"))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.name = "Cairo"
        run.font.size = Pt(16)
        run.font.bold = True

    doc.add_paragraph("")
    meta = doc.add_paragraph(
        f"الإصدار: {getattr(plan, 'version', '1.0')}  |  "
        f"التاريخ: {datetime.now().strftime('%Y-%m-%d')}  |  "
        f"الحالة: {str(getattr(plan, 'workflow_status', ''))}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.name = "Cairo"
        run.font.size = Pt(11)

    doc.add_page_break()

    # ── Document Control ──
    add_heading(doc, "١. بيانات الوثيقة", 1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    add_table_row(tbl, ["الإصدار", "التاريخ", "الوصف"], header=True)
    add_table_row(tbl, [getattr(plan, "version", "1.0"), datetime.now().strftime("%Y-%m-%d"), "الإصدار الأولي"])
    doc.add_paragraph("")

    # ── Purpose & Objectives ──
    add_heading(doc, "٢. الغرض والأهداف", 1)
    add_body(doc, getattr(plan, "purpose", "") or
             "تحدد هذه الخطة الإطار المنهجي لإدارة مخاطر المشروع وفق معايير PMI وPMBOK الإصدار الثامن.")
    doc.add_paragraph("")
    add_heading(doc, "الأهداف", 2)
    add_body(doc, getattr(plan, "objectives", "") or
             "تحديد المخاطر وتحليلها والاستجابة لها ومراقبتها بصورة منهجية طوال دورة حياة المشروع.")

    # ── Methodology ──
    add_heading(doc, "٣. المنهجية", 1)
    add_body(doc, getattr(plan, "methodology", "") or
             "تعتمد المنهجية على مزيج من التحليل النوعي (مصفوفة الاحتمالية × التأثير) "
             "والكمي (محاكاة مونت كارلو وتحليل الحساسية) لتقييم التعرض للمخاطر وترتيب أولوياتها.")

    # ── Governance ──
    add_heading(doc, "٤. الحوكمة والمسؤوليات", 1)
    add_body(doc, getattr(plan, "governance", "") or
             "يُعدّ المقاولُ سجلَ المخاطر ويُعدّ الاستشاريُّ مراجعتَه وإقرارَه، "
             "فيما يتولى ممثل المالك الموافقة النهائية وحوكمة المخاطر الاستراتيجية.")

    # ── Risk Categories ──
    add_heading(doc, "٥. فئات المخاطر المعتمدة", 1)
    categories = getattr(plan, "risk_categories", None) or [
        "مخاطر تقنية", "مخاطر تجارية", "مخاطر تعاقدية",
        "مخاطر تشغيلية", "مخاطر بيئية", "مخاطر قانونية وتنظيمية",
        "مخاطر مالية", "مخاطر أصحاب المصلحة",
    ]
    for cat in categories:
        para = doc.add_paragraph(f"• {cat}", style="List Bullet")
        set_rtl(para)
        for run in para.runs:
            run.font.name = "Cairo"
            run.font.size = Pt(11)

    # ── Probability & Impact Scales ──
    add_heading(doc, "٦. مقاييس الاحتمالية والتأثير", 1)
    add_heading(doc, "مقياس الاحتمالية", 2)
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    add_table_row(tbl2, ["المستوى", "الوصف", "النطاق المئوي", "الوزن"], header=True)
    prob_data = [
        ("1", "منخفضة جداً", "1% – 10%", "1"),
        ("2", "منخفضة", "11% – 30%", "2"),
        ("3", "متوسطة", "31% – 60%", "3"),
        ("4", "عالية", "61% – 90%", "4"),
        ("5", "عالية جداً", "91% – 100%", "5"),
    ]
    for row_data in prob_data:
        add_table_row(tbl2, row_data)

    doc.add_paragraph("")
    add_heading(doc, "أبعاد التأثير", 2)
    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = "Table Grid"
    add_table_row(tbl3, ["البُعد", "الوزن النسبي", "التعريف"], header=True)
    impact_dims = [
        ("التكلفة", "30%", "الزيادة في قيمة العقد"),
        ("الجدول الزمني", "25%", "التأخر عن التاريخ المستهدف"),
        ("النطاق", "15%", "التغير في متطلبات العمل"),
        ("الجودة", "15%", "الانحراف عن معايير القبول"),
        ("السمعة", "8%", "الأثر على صورة المنظمة"),
        ("أصحاب المصلحة", "7%", "الأثر على علاقات الأطراف"),
    ]
    for row_data in impact_dims:
        add_table_row(tbl3, row_data)

    # ── Risk Appetite & Thresholds ──
    add_heading(doc, "٧. شهية المخاطر وعتبات التصعيد", 1)
    add_body(doc, f"شهية المخاطر: {getattr(plan, 'risk_appetite', 'معتدلة') or 'معتدلة'}")
    add_body(doc, getattr(plan, "risk_thresholds", None) or
             "• المخاطر ذات درجة ≥ 15: تصعيد فوري لممثل المالك\n"
             "• المخاطر ذات درجة 9-14: إبلاغ الاستشاري ومراجعة أسبوعية\n"
             "• المخاطر ذات درجة < 9: مراقبة دورية شهرية")

    # ── Reserves ──
    add_heading(doc, "٨. الاحتياطيات المالية", 1)
    contingency = getattr(plan, "contingency_reserve_pct", 10.0)
    mgmt = getattr(plan, "management_reserve_pct", 5.0)
    add_body(doc, f"الاحتياطي الطارئ (Contingency Reserve): {contingency}% من قيمة العقد")
    add_body(doc, f"الاحتياطي الإداري (Management Reserve): {mgmt}% من قيمة العقد")

    # ── Monitoring ──
    add_heading(doc, "٩. المراقبة والإبلاغ", 1)
    add_body(doc, getattr(plan, "monitoring_approach", None) or
             "تُراجَع المخاطر ذات الأولوية العالية أسبوعياً في اجتماعات إدارة المشروع، "
             "والمخاطر المتوسطة شهرياً، وتُقدَّم تقارير المخاطر ربع السنوية لممثل المالك.")

    # ── Escalation ──
    add_heading(doc, "١٠. مسار التصعيد", 1)
    add_body(doc, getattr(plan, "escalation_path", None) or
             "المقاول (تحديد) → مدير المخاطر (تقييم) → الاستشاري (مراجعة) → "
             "ممثل المالك (قرار) → إدارة المحفظة (تصعيد استراتيجي)")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def _translate_status(status: str) -> str:
    mapping = {
        "not_started": "لم يبدأ", "in_progress": "جارٍ",
        "completed": "مكتمل", "delayed": "متأخر", "cancelled": "ملغي",
    }
    return mapping.get(status, status)


def _translate_effectiveness(eff: str) -> str:
    mapping = {
        "high": "عالية", "medium": "متوسطة",
        "low": "منخفضة", "not_assessed": "لم يُقيَّم",
    }
    return mapping.get(eff, eff)
